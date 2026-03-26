"""
Flask Web Server for VedLinks AI Educational Content Generator

NEW APPROACH:
- Uses structured topic JSON files instead of raw PDFs
- Supports lesson plan generation and question paper generation
- Topic metadata (class, subject, chapter, topics) drives all content
"""

import os
import sys
import json
import re
from pathlib import Path
from flask import Flask, render_template, jsonify, request, send_file
from werkzeug.utils import secure_filename
from flask_cors import CORS
from datetime import datetime
from io import BytesIO
import threading
import subprocess
import time
from concurrent.futures import ThreadPoolExecutor

# Fix encoding on Windows
if sys.platform == 'win32':
    os.environ['PYTHONIOENCODING'] = 'utf-8'

# Import our modules
from question_paper_generator import get_generator

app = Flask(__name__)
CORS(app)

# Configuration - NEW: Use topics directory instead of raw PDFs
DATA_TOPICS_PATH = Path("data/topics")
DATA_UPLOADS_PATH = Path("data/raw")
REGISTRY_PATH = Path("data/topic_registry.json")
FINETUNED_MODEL_PATH = Path("output/qlora_tuned_model")

# ===== Translation helpers =====
_translation_cache = {}

def translate_text(text, lang):
    """Translate a single string to target language. Returns original on error or if lang is 'en'."""
    if not text or not lang or lang == 'en':
        return text
    cache_key = f"{lang}:{text}"
    if cache_key in _translation_cache:
        return _translation_cache[cache_key]
    for attempt in range(3):  # retry up to 3 times on failure
        try:
            from deep_translator import GoogleTranslator
            result = GoogleTranslator(source='en', target=lang).translate(str(text))
            if result and result.strip():
                _translation_cache[cache_key] = result
                return result
        except Exception as e:
            if attempt < 2:
                time.sleep(0.3 * (attempt + 1))
            else:
                print(f"Translation failed for '{text[:20]}...' to '{lang}': {e}")
    return text  # return original only after all retries fail

def translate_questions_bulk(questions, lang):
    """Translate all questions in a list in parallel for maximum speed."""
    if lang == 'en' or not questions:
        return questions
    # Collect all unique strings
    all_texts = []
    for q in questions:
        for field in ['question', 'answer', 'explanation']:
            if q.get(field):
                all_texts.append(q[field])
        for opt in q.get('options', []):
            if opt:
                all_texts.append(opt)
    unique_texts = list(dict.fromkeys(all_texts))
    if not unique_texts:
        return questions
    # Translate all unique strings in parallel
    # Translate all unique strings — capped at 5 workers to avoid Google rate limits
    with ThreadPoolExecutor(max_workers=5) as ex:
        translated = list(ex.map(lambda t: translate_text(t, lang), unique_texts))
    trans_map = dict(zip(unique_texts, translated))
    # Re-assign
    for q in questions:
        for field in ['question', 'answer', 'explanation']:
            if q.get(field):
                q[field] = trans_map.get(q[field], q[field])
        if q.get('options'):
            q['options'] = [trans_map.get(opt, opt) for opt in q['options']]
    return questions

# ===== End translation helpers =====

# Global state for the fine-tuned AI model
_ai_model = None
_ai_tokenizer = None
_ai_model_loaded = False

# Global state for training orchestration
_training_status = {
    "is_training": False,
    "current_step": "idle",
    "progress": 0,
    "last_error": None,
    "start_time": None,
    "logs": []
}


def get_ai_model():
    """Lazy-load the fine-tuned LoRA model for AI question generation."""
    global _ai_model, _ai_tokenizer, _ai_model_loaded
    
    if _ai_model_loaded:
        return _ai_model, _ai_tokenizer
    
    if not FINETUNED_MODEL_PATH.exists() or not (FINETUNED_MODEL_PATH / "adapter_config.json").exists():
        print("No fine-tuned model found. AI generation will use knowledge bank only.")
        _ai_model_loaded = True
        return None, None
    
    try:
        import torch
        from peft import PeftModel, PeftConfig
        from transformers import AutoModelForCausalLM, AutoTokenizer
        
        device = "cuda" if torch.cuda.is_available() else "cpu"
        print(f"Loading fine-tuned AI model on {device}...")
        
        # Load the LoRA config to find the base model
        config = PeftConfig.from_pretrained(str(FINETUNED_MODEL_PATH))
        base_model_name = config.base_model_name_or_path
        
        # Load tokenizer
        _ai_tokenizer = AutoTokenizer.from_pretrained(str(FINETUNED_MODEL_PATH))
        if _ai_tokenizer.pad_token is None:
            _ai_tokenizer.pad_token = _ai_tokenizer.eos_token
        
        # Load base model
        device = "cuda" if torch.cuda.is_available() else "cpu"
        if device == "cuda":
            base_model = AutoModelForCausalLM.from_pretrained(
                base_model_name,
                torch_dtype=torch.float16,
                device_map="auto",
                trust_remote_code=True,
            )
        else:
            base_model = AutoModelForCausalLM.from_pretrained(
                base_model_name,
                trust_remote_code=True,
            )
        
        # Load LoRA adapter
        _ai_model = PeftModel.from_pretrained(base_model, str(FINETUNED_MODEL_PATH))
        _ai_model.eval()
        
        _ai_model_loaded = True
        print(f"AI model loaded on {device}")
        return _ai_model, _ai_tokenizer
        
    except Exception as e:
        print(f"Failed to load AI model: {e}")
        _ai_model_loaded = True
        return None, None


def generate_ai_text(prompt_text, max_new_tokens=200):
    """Generate text using the fine-tuned model or HF Inference API."""
    import os
    use_hf_api = os.environ.get('USE_HF_API', 'False').lower() == 'true'
    
    if use_hf_api:
        import requests
        hf_token = os.environ.get('HF_API_TOKEN', '')
        hf_model = os.environ.get('HF_MODEL_ID', '')
        hf_space_url = os.environ.get('HF_SPACE_URL', '')
        
        if not hf_token and not hf_space_url:
            print("HF_API_TOKEN or HF_SPACE_URL is required.")
            return None
            
        headers = {}
        if hf_token:
            headers["Authorization"] = f"Bearer {hf_token}"
            
        payload = {
            "inputs": prompt_text,
            "parameters": {
                "max_new_tokens": max_new_tokens,
                "temperature": 0.7,
                "top_p": 0.9,
                "repetition_penalty": 1.2,
                "return_full_text": False
            }
        }
        
        # Determine the API URL: prioritize Space URL if provided
        if hf_space_url:
            API_URL = hf_space_url
        elif hf_model:
            API_URL = f"https://api-inference.huggingface.co/models/{hf_model}"
        else:
            print("Neither HF_SPACE_URL nor HF_MODEL_ID were provided.")
            return None
            
        try:
            response = requests.post(API_URL, headers=headers, json=payload, timeout=30)
            if response.status_code == 200:
                result = response.json()
                if isinstance(result, list) and len(result) > 0 and 'generated_text' in result[0]:
                    generated = result[0]['generated_text']
                    if "### Response:" in generated:
                        response_text = generated.split("### Response:")[-1].strip()
                    else:
                        response_text = generated.strip()
                    return response_text if response_text else None
            print(f"HF API Error: {response.status_code} - {response.text}")
            return None
        except Exception as e:
            print(f"HF API Exception: {e}")
            return None
            
    # Local generation fallback
    import torch
    model, tokenizer = get_ai_model()
    if model is None or tokenizer is None:
        return None
    
    try:
        inputs = tokenizer(prompt_text, return_tensors="pt", truncation=True, max_length=256)
        device = next(model.parameters()).device
        inputs = {k: v.to(device) for k, v in inputs.items()}
        
        with torch.no_grad():
            outputs = model.generate(
                **inputs,
                max_new_tokens=max_new_tokens,
                temperature=0.7,
                top_p=0.9,
                do_sample=True,
                pad_token_id=tokenizer.eos_token_id,
                repetition_penalty=1.2,
            )
        
        generated = tokenizer.decode(outputs[0], skip_special_tokens=True)
        if "### Response:" in generated:
            response = generated.split("### Response:")[-1].strip()
        else:
            response = generated[len(prompt_text):].strip()
        
        return response if response else None
    except Exception as e:
        print(f"AI generation error: {e}")
        return None


def load_topic_registry() -> dict:
    """Load the topic registry for fallback topic data."""
    if REGISTRY_PATH.exists():
        try:
            with open(REGISTRY_PATH, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"Error loading registry: {e}")
    return {"files": {}}


def get_available_topics():
    """Get list of available topics from data/topics directory and registry.
    
    Topics can come from:
    1. Individual JSON files in data/topics/
    2. Topic registry (data/topic_registry.json) for entries without JSON files
    """
    topics = []
    seen_chapters = set()
    
    # 1. Load from individual topic JSON files (primary source)
    if DATA_TOPICS_PATH.exists():
        for file_path in DATA_TOPICS_PATH.glob("*.json"):
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    topic_data = json.load(f)
                    
                chapter = topic_data.get('chapter', file_path.stem)
                seen_chapters.add(chapter)
                
                topics.append({
                    'id': file_path.name,
                    'name': chapter,
                    'class': topic_data.get('class', 'N/A'),
                    'subject': topic_data.get('subject', 'N/A'),
                    'chapter': chapter,
                    'chapter_number': topic_data.get('chapter_number', 0),
                    'topics': topic_data.get('topics', []),
                    'topic_count': len(topic_data.get('topics', [])),
                    'source': 'topic_file'
                })
            except Exception as e:
                print(f"Error loading {file_path.name}: {e}")
    
    # 2. Load from registry for any entries not already covered
    registry = load_topic_registry()
    for file_id, entry in registry.get('files', {}).items():
        chapter = entry.get('chapter', '')
        
        # Skip if we already have this chapter from a topic file
        if chapter in seen_chapters:
            continue
        
        # Skip empty/unknown chapters
        if not chapter or chapter.startswith('Unknown'):
            continue
        
        topics.append({
            'id': file_id,
            'name': chapter,
            'class': entry.get('class', 'N/A'),
            'subject': entry.get('subject', 'N/A'),
            'chapter': chapter,
            'chapter_number': entry.get('chapter_number', 0),
            'topics': entry.get('topics', []),
            'topic_count': len(entry.get('topics', [])),
            'source': 'registry'
        })
        seen_chapters.add(chapter)
    
    return sorted(topics, key=lambda x: (x['class'], x['subject'], x['name']))


def load_topic_data(topic_id: str) -> dict:
    """Load complete topic data from a JSON file or registry.
    
    First checks for individual topic file in data/topics/,
    then falls back to the topic registry.
    """
    # 1. Try loading from individual topic file
    file_path = DATA_TOPICS_PATH / topic_id
    
    if file_path.exists():
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"Error loading {topic_id}: {e}")
    
    # 2. Try loading from registry
    registry = load_topic_registry()
    if topic_id in registry.get('files', {}):
        entry = registry['files'][topic_id]
        return {
            'class': entry.get('class', 'N/A'),
            'subject': entry.get('subject', 'N/A'),
            'chapter': entry.get('chapter', 'Unknown'),
            'chapter_number': entry.get('chapter_number', 0),
            'topics': entry.get('topics', [])
        }
    
    return {}


@app.route('/')
def index():
    """Serve the main question paper generator page."""
    return render_template('index.html')


@app.route('/practice')
def practice():
    """Serve the practice and doubts page for students."""
    return render_template('practice.html')


@app.route('/lesson-planner')
def lesson_planner():
    """Serve the lesson plan generator page."""
    return render_template('lesson_planner.html')


@app.route('/upload')
def upload_page():
    """Serve the textbook upload page."""
    return render_template('upload.html')


@app.route('/api/upload-textbook', methods=['POST'])
def api_upload_textbook():
    """Upload one or more textbook PDFs with manual metadata."""
    try:
        # Get the uploaded files (supports multiple)
        files = request.files.getlist('files')
        
        # Fallback: also check for single 'file' key (backward compatibility)
        if not files:
            if 'file' in request.files:
                files = [request.files['file']]
            else:
                return jsonify({'success': False, 'error': 'No files uploaded'}), 400
        
        # Filter out empty filenames
        files = [f for f in files if f.filename and f.filename != '']
        if not files:
            return jsonify({'success': False, 'error': 'No files selected'}), 400
        
        # Validate all files are PDFs
        for f in files:
            if not f.filename.lower().endswith('.pdf'):
                return jsonify({'success': False, 'error': f'Only PDF files are allowed. "{f.filename}" is not a PDF.'}), 400
        
        # Get manual metadata from form
        class_num = request.form.get('class', '').strip()
        subject = request.form.get('subject', '').strip()
        chapter_number = request.form.get('chapter_number', '').strip()
        chapter_name = request.form.get('chapter_name', '').strip()
        topics_raw = request.form.get('topics', '').strip()
        
        # Validate required fields
        if not class_num:
            return jsonify({'success': False, 'error': 'Class is required'}), 400
        if not subject:
            return jsonify({'success': False, 'error': 'Subject is required'}), 400
        if not chapter_number:
            return jsonify({'success': False, 'error': 'Chapter number is required'}), 400
        if not chapter_name:
            return jsonify({'success': False, 'error': 'Chapter name is required'}), 400
        
        chapter_number = int(chapter_number)
        
        # Parse topics (comma-separated)
        topics = [t.strip() for t in topics_raw.split(',') if t.strip()] if topics_raw else []
        
        # Save each PDF and create topic entries
        raw_dir = Path('data/raw')
        raw_dir.mkdir(parents=True, exist_ok=True)
        topics_dir = Path('data/topics')
        topics_dir.mkdir(parents=True, exist_ok=True)
        
        saved_files = []
        extracted_texts = []
        registry = load_topic_registry()
        
        # Create extracted text directory
        extracted_dir = Path('data/extracted')
        extracted_dir.mkdir(parents=True, exist_ok=True)
        
        for file in files:
            safe_name = secure_filename(file.filename)
            pdf_path = raw_dir / safe_name
            file.save(str(pdf_path))
            saved_files.append(safe_name)
            
            # Extract text from PDF immediately for question generation
            try:
                from src.pdf_processor import get_chapter_text
                text_content = get_chapter_text(str(pdf_path))
                if text_content and len(text_content.strip()) > 50:
                    text_file = extracted_dir / f"{safe_name}.txt"
                    with open(text_file, 'w', encoding='utf-8') as tf:
                        tf.write(text_content)
                    extracted_texts.append(str(text_file))
                    print(f"  Extracted {len(text_content)} chars from {safe_name}")
                else:
                    print(f"  Warning: No meaningful text extracted from {safe_name}")
            except Exception as e:
                print(f"  Error extracting text from {safe_name}: {e}")
        
        # Create topic JSON filename
        subject_slug = re.sub(r'[^a-z0-9]', '_', subject.lower()).strip('_')
        topic_filename = f"class_{class_num}_{subject_slug}_ch{chapter_number}.json"
        
        # Create topic data
        topic_data = {
            'class': str(class_num),
            'subject': subject,
            'chapter': chapter_name,
            'chapter_number': chapter_number,
            'topics': topics,
            'source_pdfs': saved_files
        }
        
        # Save to data/topics/
        topic_path = topics_dir / topic_filename
        
        with open(topic_path, 'w', encoding='utf-8') as f:
            json.dump(topic_data, f, indent=2, ensure_ascii=False)
        
        # Update registry
        registry['files'][topic_filename] = {
            'class': str(class_num),
            'subject': subject,
            'chapter_number': chapter_number,
            'chapter': chapter_name,
            'topics': topics,
            'has_knowledge_bank': False,
            'source_pdfs': saved_files
        }
        
        with open(REGISTRY_PATH, 'w', encoding='utf-8') as f:
            json.dump(registry, f, indent=2, ensure_ascii=False)
        
        file_count = len(saved_files)
        return jsonify({
            'success': True,
            'message': f'Uploaded {file_count} file{"s" if file_count > 1 else ""} successfully! Class {class_num} {subject} - Ch {chapter_number}: {chapter_name}',
            'topic_file': topic_filename,
            'pdf_files': saved_files
        })
        
    except ValueError:
        return jsonify({'success': False, 'error': 'Chapter number must be a valid number'}), 400
    except Exception as e:
        print(f'Error uploading textbook: {e}')
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/topics', methods=['GET'])
def api_get_topics():
    """Get list of available topics."""
    topics = get_available_topics()
    return jsonify({
        'success': True,
        'topics': topics,
        'count': len(topics)
    })


@app.route('/api/topics-grouped', methods=['GET'])
def api_get_topics_grouped():
    """Get topics grouped hierarchically: Class -> Subject -> Chapters."""
    topics = get_available_topics()
    
    grouped = {}
    for t in topics:
        cls = t.get('class', 'N/A')
        subject = t.get('subject', 'N/A')
        
        if cls not in grouped:
            grouped[cls] = {}
        if subject not in grouped[cls]:
            grouped[cls][subject] = []
        
        grouped[cls][subject].append({
            'id': t['id'],
            'chapter': t.get('chapter', t.get('name', '')),
            'chapter_number': t.get('chapter_number', 0),
            'topic_count': t.get('topic_count', 0),
            'topics': t.get('topics', []),
            'source': t.get('source', 'unknown')
        })
    
    # Sort chapters by chapter_number within each subject
    for cls in grouped:
        for subject in grouped[cls]:
            grouped[cls][subject].sort(key=lambda x: x.get('chapter_number', 0))
    
    # Sort classes numerically
    sorted_grouped = {}
    for cls in sorted(grouped.keys(), key=lambda x: int(x) if x.isdigit() else 999):
        sorted_grouped[cls] = grouped[cls]
    
    return jsonify({
        'success': True,
        'grouped': sorted_grouped
    })


@app.route('/api/topic/<topic_id>', methods=['GET'])
def api_get_topic_details(topic_id):
    """Get detailed information about a specific topic."""
    topic_data = load_topic_data(topic_id)
    
    if not topic_data:
        return jsonify({
            'success': False,
            'error': f'Topic not found: {topic_id}'
        }), 404
    
    return jsonify({
        'success': True,
        'topic': topic_data
    })


@app.route('/api/practice-questions', methods=['POST'])
def api_practice_questions():
    """Get practice questions for a topic, mixing knowledge bank and AI-generated."""
    import random
    from question_paper_generator import NCERT_KNOWLEDGE, get_generator
    
    data = request.get_json()
    topic_id = data.get('topicId', '')
    lang = data.get('lang', 'en')

    # Load topic data
    topic_data = load_topic_data(topic_id)
    if not topic_data:
        return jsonify({'success': False, 'error': 'Topic not found'}), 404

    chapter = topic_data.get('chapter', '')

    # Use generator's fuzzy matching instead of exact dict lookup
    generator = get_generator()
    content_str = f"Chapter: {chapter}\nSourcePDFs: {', '.join(topic_data.get('source_pdfs', []))}"
    knowledge = generator.get_chapter_knowledge(content_str) or {}
    
    # Add MCQs from KB
    for item in knowledge.get('mcq_pool', []):
        q = item[0]
        opts = item[1]
        ans = item[2]
        exp = item[3] if len(item) > 3 else ""
        img = item[4] if len(item) > 4 else None
        
        q_obj = {
            'type': 'mcq',
            'question': q,
            'options': opts,
            'answer': ans,
            'explanation': exp,
            'source': 'knowledge_bank'
        }
        if img: q_obj['imageUrl'] = img
        questions.append(q_obj)
    
    # Add fill in blanks from KB
    for q, ans in knowledge.get('fill_blanks', []):
        questions.append({
            'type': 'fill_blank',
            'question': q.replace('_______', '________'),
            'answer': ans,
            'source': 'knowledge_bank'
        })
    
    # Add short answer questions from KB
    for item in knowledge.get('short_answers', []):
        q = item[0]
        ans = item[1]
        img = item[2] if len(item) > 2 else None
        
        q_obj = {
            'type': 'short_answer',
            'question': q,
            'answer': ans,
            'source': 'knowledge_bank'
        }
        if img: q_obj['imageUrl'] = img
        questions.append(q_obj)
    
    # Check if AI model is available
    use_hf_api = os.environ.get('USE_HF_API', 'False').lower() == 'true'
    has_ai = use_hf_api or (FINETUNED_MODEL_PATH.exists() and (FINETUNED_MODEL_PATH / "adapter_config.json").exists())
    
    # DYNAMIC FALLBACK: If knowledge bank is empty, generate AI questions
    ai_questions = []
    if not questions and has_ai:
        try:
            print(f"No KB content for '{chapter}'. Generating dynamic AI questions...")
            for _ in range(5): # Generate 5 initial AI questions
                q_type = random.choice(['mcq', 'short_answer'])
                prompt = f"### Instruction:\nGenerate a {q_type} question about '{chapter}'.\n\n### Response:"
                
                response = generate_ai_text(prompt, max_new_tokens=250)
                if response:
                    if q_type == 'mcq' and 'Question:' in response:
                        lines = response.strip().split('\n')
                        q_text, opts, ans, exp = '', [], '', ''
                        for line in lines:
                            line = line.strip()
                            if line.startswith('Question:'): q_text = line.replace('Question:', '').strip()
                            elif line.startswith(('A)', 'B)', 'C)', 'D)')): opts.append(line[2:].strip())
                            elif line.startswith('Answer:'): ans = line.replace('Answer:', '').strip()
                            elif line.startswith('Explanation:'): exp = line.replace('Explanation:', '').strip()
                        
                        if q_text and len(opts) >= 4 and ans:
                            ai_questions.append({
                                'type': 'mcq', 'question': q_text, 'options': opts[:4],
                                'answer': ans, 'explanation': exp, 'source': 'ai_generated'
                            })
                    elif q_type == 'short_answer' and 'Question:' in response:
                        parts = response.split('Answer:')
                        q_text = parts[0].replace('Question:', '').strip()
                        ans_text = parts[1].strip() if len(parts) > 1 else ''
                        if q_text and ans_text:
                            ai_questions.append({
                                'type': 'short_answer', 'question': q_text, 
                                'answer': ans_text, 'source': 'ai_generated'
                            })
        except Exception as e:
            print(f"Error generating fallback AI questions: {e}")
    
    # PDF TEXT FALLBACK: If still no questions, extract from PDF text
    if not questions and not ai_questions:
        try:
            pdf_questions = generator._generate_fallback_questions('mcq', 1, 5, content_str)
            pdf_questions += generator._generate_fallback_questions('short', 2, 5, content_str)
            for pq in pdf_questions:
                pq['source'] = 'pdf_extracted'
            ai_questions.extend(pdf_questions)
            print(f"Generated {len(pdf_questions)} questions from PDF text for '{chapter}'")
        except Exception as e:
            print(f"Error generating PDF-based practice questions: {e}")

    # Combine and Shuffle
    random.shuffle(questions)
    final_questions = (questions + ai_questions)[:15]
    random.shuffle(final_questions)

    # Translate if needed
    translate_questions_bulk(final_questions, lang)

    return jsonify({
        'success': True,
        'questions': final_questions,
        'count': len(final_questions),
        'ai_available': has_ai,
        'chapter': chapter
    })


@app.route('/api/concepts', methods=['POST'])
def api_concepts():
    """Get key concepts for a topic."""
    from question_paper_generator import NCERT_KNOWLEDGE, get_generator
    
    data = request.get_json()
    topic_id = data.get('topicId', '')
    lang = data.get('lang', 'en')

    # Load topic data
    topic_data = load_topic_data(topic_id)
    if not topic_data:
        return jsonify({'success': False, 'error': 'Topic not found'}), 404

    chapter = topic_data.get('chapter', '')

    # Use fuzzy matching for knowledge bank lookup
    generator = get_generator()
    content_str = f"Chapter: {chapter}\nSourcePDFs: {', '.join(topic_data.get('source_pdfs', []))}"
    knowledge = generator.get_chapter_knowledge(content_str) or {}
    
    # Get concepts from knowledge bank
    concepts = []
    
    # Add key concepts
    for term, definition in knowledge.get('key_concepts', []):
        concepts.append({
            'term': term,
            'definition': definition,
            'source': 'knowledge_bank'
        })
    
    # Also add short answers as additional Q&A concepts
    for item in knowledge.get('short_answers', []):
        q = item[0]
        ans = item[1]
        concepts.append({
            'term': q,
            'definition': ans,
            'source': 'knowledge_bank'
        })
    
    # Add from topics list if still empty
    if not concepts and topic_data.get('topics'):
        for topic in topic_data['topics']:
            concepts.append({
                'term': topic,
                'definition': f'Key topic in {chapter}. Review your textbook for detailed explanation.',
                'source': 'topic_list'
            })
            
    # AI DYNAMIC FALLBACK: Add AI-generated insight if knowledge bank is thin
    use_hf_api = os.environ.get('USE_HF_API', 'False').lower() == 'true'
    has_ai = use_hf_api or (FINETUNED_MODEL_PATH.exists() and (FINETUNED_MODEL_PATH / "adapter_config.json").exists())
    if has_ai:
        try:
            # Determine subtopics to generate insights for
            subtopics_to_generate = topic_data.get('topics', [])
            
            # If no formal topics list exists, try to extract unique terms from the knowledge bank
            kb_concepts = knowledge.get('key_concepts', [])
            if not subtopics_to_generate and len(kb_concepts) > 0:
                subtopics_to_generate = [term for term, _ in kb_concepts]
                
            # If STILL no subtopics (empty KB), extract headings directly from the source PDF
            if not subtopics_to_generate and topic_data.get('source_pdfs'):
                from src.pdf_processor import get_chapter_text
                safe_chap = str(chapter).encode('ascii', 'replace').decode('ascii')
                print(f"No known subtopics for '{safe_chap}'. Extracting from PDF...")
                
                # Take the first associated PDF for this chapter
                pdf_filename = topic_data['source_pdfs'][0]
                pdf_path = DATA_UPLOADS_PATH / pdf_filename
                
                if pdf_path.exists():
                    raw_text = get_chapter_text(str(pdf_path))
                    # Find all headings (lines starting with ## from our parser)
                    import re
                    headings = re.findall(r'##\s+([^\n]+)', raw_text)
                    
                    if headings:
                        # Clean and filter headings
                        cleaned_headings = []
                        seen_headings = set()
                        for h in headings:
                            clean = h.strip()
                            clean_upper = clean.upper()
                            # Extract only valid numbered subtopics (e.g. "7.1", "7.1.1 Fission")
                            if 3 < len(clean) < 100 and clean not in seen_headings:
                                # Look for a number pattern like X.Y or X.Y.Z
                                has_number = re.match(r'^([0-9]+\.[0-9]+(\.[0-9]+)*)\s*(.*)', clean)
                                if has_number:
                                    # Strip the number prefix to get the clean topic text
                                    topic_text = has_number.group(3).strip()
                                    
                                    # General text cleanup for repeating chunks like "MA MA MA" because of poor PDF parsing
                                    topic_text = re.sub(r'([A-Z])([\s-]+\1)+', r'\1', topic_text)
                                    topic_text = re.sub(r'\?+', '?', topic_text)
                                    topic_text = topic_text.strip('- ')
                                    
                                    if len(topic_text) > 3 and topic_text.upper() != 'QUESTION' and topic_text.upper() != 'QUESTIONS':
                                        cleaned_headings.append(topic_text)
                                        seen_headings.add(topic_text)
                        
                        subtopics_to_generate = cleaned_headings[:5]
                        print(f"PARSED SUBTOPICS FROM PDF: {subtopics_to_generate}")
                else:
                    print(f"Source PDF not found: {pdf_path}")
            
            # Generate a summary for each subtopic if we have few concepts overall
            if len(concepts) < 3 and subtopics_to_generate:
                safe_chap = str(chapter).encode('ascii', 'replace').decode('ascii')
                print(f"Few concepts for '{safe_chap}'. Generating AI overviews for subtopics: {subtopics_to_generate[:5]}")
                # Remove the generic "Review your textbook" placeholders we added earlier
                concepts = [c for c in concepts if c.get('source') != 'topic_list']
                
                # We cap at 5 topics to avoid slow generation times
                for idx, topic in enumerate(subtopics_to_generate[:5]):
                    safe_topic = str(topic).encode('ascii', 'replace').decode('ascii')
                    print(f"Generating AI concept for subtopic: {safe_topic}...")
                    prompt = f"### Instruction:\nExplain the core scientific concepts of '{topic}' in the context of the chapter '{chapter}' in a detailed paragraph as an expert teacher.\n\n### Response:"
                    ai_insight = generate_ai_text(prompt, max_new_tokens=300)
                    if ai_insight:
                        concepts.append({
                            'term': topic,
                            'definition': ai_insight,
                            'source': 'ai_generated'
                        })
                        
            # If still empty (no subtopics found), fallback to single overview
            if len(concepts) == 0 and not subtopics_to_generate:
                safe_chap = str(chapter).encode('ascii', 'replace').decode('ascii')
                print(f"No subtopics found for '{safe_chap}'. Generating generic AI overview...")
                prompt = f"### Instruction:\nExplain the core scientific concepts of '{chapter}' in a detailed paragraph as an expert teacher.\n\n### Response:"
                ai_insight = generate_ai_text(prompt, max_new_tokens=350)
                if ai_insight:
                    concepts.insert(0, {
                        'term': 'AI Topic Overview',
                        'definition': ai_insight,
                        'source': 'ai_generated'
                    })
        except Exception as e:
            print(f"Error generating AI concept insight: {e}")

    # Translate concepts if needed
    if lang != 'en':
        with ThreadPoolExecutor(max_workers=5) as ex:
            term_futs = [ex.submit(translate_text, c.get('term', ''), lang) for c in concepts]
            def_futs = [ex.submit(translate_text, c.get('definition', ''), lang) for c in concepts]
        for i, c in enumerate(concepts):
            c['term'] = term_futs[i].result()
            c['definition'] = def_futs[i].result()

    return jsonify({
        'success': True,
        'concepts': concepts,
        'count': len(concepts),
        'ai_available': has_ai
    })


@app.route('/api/generate-paper', methods=['POST'])
def api_generate_paper():
    """Generate a complete question paper."""
    try:
        config = request.get_json()
        
        if not config:
            return jsonify({
                'success': False,
                'error': 'No configuration provided'
            }), 400
        
        # Validate required fields
        if 'sections' not in config or not config['sections']:
            return jsonify({
                'success': False,
                'error': 'At least one section is required'
            }), 400
        
        if 'selectedTopics' not in config or not config['selectedTopics']:
            return jsonify({
                'success': False,
                'error': 'At least one topic must be selected'
            }), 400
        
        # Get language for translation
        lang = config.get('lang', 'en')

        # Load topic data (NEW: Uses structured topic data)
        topic_contents = {}
        topic_metadata = []
        
        for topic_id in config['selectedTopics']:
            topic_data = load_topic_data(topic_id)
            if topic_data:
                # Create content from topic metadata for the generator
                content = f"Class: {topic_data.get('class', '')}\n"
                content += f"Subject: {topic_data.get('subject', '')}\n"
                content += f"Chapter: {topic_data.get('chapter', '')}\n"
                content += f"Topics: {', '.join(topic_data.get('topics', []))}\n"
                source_pdfs = topic_data.get('source_pdfs', [])
                if source_pdfs:
                    content += f"SourcePDFs: {', '.join(source_pdfs)}"
                
                topic_contents[topic_id] = content
                topic_metadata.append(topic_data)
        
        if not topic_contents:
            return jsonify({
                'success': False,
                'error': 'Could not load any topic data'
            }), 400
        
        # Get generator and generate paper
        generator = get_generator()
        paper = generator.generate_paper(config, topic_contents)
        
        # Add metadata about topics used
        paper['topicMetadata'] = topic_metadata

        # Translate AI-generated content if language is not English
        if lang != 'en':
            for section in paper.get('sections', []):
                translate_questions_bulk(section.get('questions', []), lang)
            # Translate answer key answers
            for section_key in paper.get('answerKey', []):
                for ans in section_key.get('answers', []):
                    if ans.get('answer'):
                        ans['answer'] = translate_text(ans['answer'], lang)
                    if ans.get('explanation'):
                        ans['explanation'] = translate_text(ans['explanation'], lang)

        return jsonify({
            'success': True,
            'paper': paper
        })
        
    except Exception as e:
        print(f"Error generating paper: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/generate-lesson-plan', methods=['POST'])
def api_generate_lesson_plan():
    """Generate a detailed school-format lesson plan."""
    try:
        from question_paper_generator import NCERT_KNOWLEDGE
        
        config = request.get_json()

        if not config:
            return jsonify({
                'success': False,
                'error': 'No configuration provided'
            }), 400

        lang = config.get('lang', 'en')

        # Get chapter info
        chapter_id = config.get('chapterId')
        chapter_name = config.get('chapterName', 'Chapter')
        
        # Load topic data
        topic_data = load_topic_data(chapter_id) if chapter_id else {}
        
        # Get knowledge from NCERT bank
        knowledge = NCERT_KNOWLEDGE.get(chapter_name, {})
        key_concepts = knowledge.get('key_concepts', [])
        
        # Calculate timing based on periods
        periods = int(config.get('periodsCount', 1))
        duration = int(config.get('periodDuration', 40))
        total_time = periods * duration
        
        # Distribute time: 15% opening, 70% main, 15% closure
        opening_time = max(5, int(total_time * 0.15))
        main_time = int(total_time * 0.70)
        closure_time = max(5, int(total_time * 0.15))
        
        # Build prerequisite knowledge from key concepts
        prereq = []
        if key_concepts:
            for i, (term, defn) in enumerate(key_concepts[:4]):
                prereq.append(f"{term}: {defn[:80]}..." if len(defn) > 80 else f"{term}: {defn}")
        else:
            prereq = [
                "Basic understanding of the subject",
                "Previous chapter concepts",
                "Fundamental terminology"
            ]
        
        # Build teaching phases
        phases = [
            {
                "name": "Opening / Activating Prior Knowledge",
                "duration": f"{opening_time} min",
                "learningOutcome": f"Students will recall previous knowledge related to {config.get('topic', chapter_name)}",
                "methodology": ["Lecture", "Discussion", "Q&A", "Brainstorming"],
                "teachingAids": ["Blackboard", "Previous notes"],
                "teacherActivities": [
                    "Recap previous lesson",
                    "Ask probing questions to assess prior knowledge",
                    "Introduce today's topic with real-life example"
                ],
                "learnerActivities": [
                    "Answer pre-requisite questions",
                    "Share previous knowledge",
                    "Listen to introduction"
                ],
                "assessment": "Oral questioning to check prerequisite knowledge",
                "homeAssignment": ""
            },
            {
                "name": "Teaching & Conducting Activities",
                "duration": f"{main_time} min",
                "learningOutcome": f"Students will understand and apply concepts of {config.get('topic', chapter_name)}",
                "methodology": [
                    "Active Collaborative Learning",
                    "Inquiry-Based Learning",
                    "Experiential Learning",
                    "Concept Mapping"
                ],
                "teachingAids": ["PPT", "Videos", "Charts", "Models", "Diagrams"],
                "teacherActivities": [
                    "Explain key concepts with examples",
                    "Show PPT/Video demonstration",
                    "Conduct hands-on activity",
                    "Facilitate group discussion",
                    "Address misconceptions"
                ],
                "learnerActivities": [
                    "Take notes and draw diagrams",
                    "Observe demonstrations",
                    "Participate in activities",
                    "Ask questions and clarify doubts",
                    "Work in groups on assigned tasks"
                ],
                "assessment": "Formative assessment through observation and questioning",
                "homeAssignment": ""
            },
            {
                "name": "Closure, Feedback & Doubt Clarification",
                "duration": f"{closure_time} min",
                "learningOutcome": "Students will consolidate learning and clarify any remaining doubts",
                "methodology": ["Summarization", "Peer Teaching", "Reflection"],
                "teachingAids": ["Worksheet", "Blackboard"],
                "teacherActivities": [
                    "Summarize key points",
                    "Conduct quick assessment",
                    "Address remaining doubts",
                    "Assign homework"
                ],
                "learnerActivities": [
                    "Participate in recap discussion",
                    "Complete quick assessment",
                    "Ask clarifying questions",
                    "Note down homework"
                ],
                "assessment": "Quick quiz or exit ticket",
                "homeAssignment": f"Read textbook pages on {config.get('topic', chapter_name)}. Complete worksheet questions 1-5."
            }
        ]
        
        # Build values based on chapter
        values_map = {
            "Chemical Reactions": "Scientific inquiry, Environmental awareness, Safety consciousness",
            "Acids": "Curiosity, Analytical thinking, Application of knowledge",
            "Metals": "Resource conservation, Sustainable development, Scientific temperament",
            "Carbon": "Environmental responsibility, Organic awareness, Practical application",
            "Life Processes": "Respect for life, Health consciousness, Scientific observation",
            "Control": "Self-discipline, Mind-body awareness, Coordination",
            "Reproduction": "Respect for life, Responsibility, Scientific understanding",
            "Heredity": "Appreciation of diversity, Scientific reasoning, Family values",
            "Light": "Scientific curiosity, Observation skills, Optical awareness",
            "Human Eye": "Health awareness, Vision care, Scientific appreciation",
            "Electricity": "Energy conservation, Safety awareness, Practical application",
            "Magnetic": "Scientific inquiry, Technological appreciation, Practical skills",
            "Environment": "Environmental stewardship, Conservation, Sustainable living"
        }
        
        # Find matching value
        values = "Scientific inquiry, Critical thinking, Curiosity"
        for key, val in values_map.items():
            if key.lower() in chapter_name.lower():
                values = val
                break
        
        # Build real-life applications
        applications_map = {
            "Chemical Reactions": "Cooking, rusting prevention, fire extinguishers, photography",
            "Acids": "Antacids for digestion, cleaning agents, food preservation, batteries",
            "Metals": "Construction, jewelry, electrical wiring, utensils, coins",
            "Carbon": "Fuels, plastics, medicines, soaps, cosmetics",
            "Life Processes": "Nutrition, breathing exercises, blood donation awareness",
            "Control": "Stress management, reflex actions, hormonal health",
            "Reproduction": "Agriculture, animal husbandry, family planning awareness",
            "Heredity": "Genetic counseling, crop improvement, disease prediction",
            "Light": "Photography, fiber optics, rainbows, mirrors",
            "Human Eye": "Spectacles, cameras, telescopes, corrective surgery",
            "Electricity": "Home appliances, lighting, electronic devices, power generation",
            "Magnetic": "Motors, generators, MRI scans, magnetic storage",
            "Environment": "Waste management, conservation, sustainable practices"
        }
        
        real_life = "Everyday science applications and practical demonstrations"
        for key, val in applications_map.items():
            if key.lower() in chapter_name.lower():
                real_life = val
                break
        
        # Build cross-curricular connections
        cross_curricular = "Mathematics (calculations), Geography (resources), History (scientific discoveries)"
        
        # Build lesson plan response
        lesson_plan = {
            "teacherName": config.get('teacherName', 'Teacher'),
            "schoolName": config.get('schoolName', 'School Name'),
            "designation": config.get('designation', ''),
            "grade": config.get('grade', '10'),
            "subject": config.get('subject', 'Science'),
            "chapterName": chapter_name,
            "topic": config.get('topic', chapter_name),
            "date": config.get('date', ''),
            "planNumber": config.get('planNumber', '1'),
            "periodsCount": periods,
            "periodDuration": duration,
            "totalDuration": total_time,
            "prerequisiteKnowledge": prereq,
            "phases": phases,
            "values": values,
            "realLifeApplication": real_life,
            "crossCurricular": cross_curricular,
            "extendedTask": f"Research project: Explore real-world applications of {config.get('topic', chapter_name)} and prepare a short presentation."
        }
        
        # Translate lesson plan content if needed
        if lang != 'en':
            # Build flat list of (path, value) for all translatable strings
            # so we can translate all in parallel and write back easily
            jobs = []  # list of (setter_fn, text)

            def _add(setter, text):
                jobs.append((setter, text or ''))

            _add(lambda v: lesson_plan.__setitem__('values', v), lesson_plan.get('values', ''))
            _add(lambda v: lesson_plan.__setitem__('realLifeApplication', v), lesson_plan.get('realLifeApplication', ''))
            _add(lambda v: lesson_plan.__setitem__('crossCurricular', v), lesson_plan.get('crossCurricular', ''))
            _add(lambda v: lesson_plan.__setitem__('extendedTask', v), lesson_plan.get('extendedTask', ''))

            prereqs = lesson_plan.get('prerequisiteKnowledge', [])
            for i, p in enumerate(prereqs):
                idx_capture = i
                _add(lambda v, i=idx_capture: prereqs.__setitem__(i, v), p)

            for phase in lesson_plan.get('phases', []):
                ph = phase  # capture
                _add(lambda v, p=ph: p.__setitem__('name', v), phase.get('name', ''))
                _add(lambda v, p=ph: p.__setitem__('learningOutcome', v), phase.get('learningOutcome', ''))
                _add(lambda v, p=ph: p.__setitem__('assessment', v), phase.get('assessment', ''))
                _add(lambda v, p=ph: p.__setitem__('homeAssignment', v), phase.get('homeAssignment', ''))
                for i, act in enumerate(phase.get('teacherActivities', [])):
                    acts = phase['teacherActivities']
                    _add(lambda v, a=acts, i=i: a.__setitem__(i, v), act)
                for i, act in enumerate(phase.get('learnerActivities', [])):
                    acts = phase['learnerActivities']
                    _add(lambda v, a=acts, i=i: a.__setitem__(i, v), act)
                for i, m in enumerate(phase.get('methodology', [])):
                    mlist = phase['methodology']
                    _add(lambda v, ml=mlist, i=i: ml.__setitem__(i, v), m)
                for i, aid in enumerate(phase.get('teachingAids', [])):
                    alist = phase['teachingAids']
                    _add(lambda v, al=alist, i=i: al.__setitem__(i, v), aid)

            # Translate all strings — capped at 5 workers to avoid Google rate limits
            with ThreadPoolExecutor(max_workers=5) as ex:
                futs = [ex.submit(translate_text, text, lang) for _, text in jobs]
            for (setter, _), fut in zip(jobs, futs):
                setter(fut.result())

        return jsonify({
            'success': True,
            'lessonPlan': lesson_plan
        })
        
    except Exception as e:
        print(f"Error generating lesson plan: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/model-status', methods=['GET'])
def api_model_status():
    """Check if model is loaded."""
    generator = get_generator()
    return jsonify({
        'success': True,
        'loaded': generator.is_loaded,
        'device': generator.device or 'not loaded',
        'mode': 'template' if generator.use_template_mode else 'model'
    })


@app.route('/api/load-model', methods=['POST'])
def api_load_model():
    """Preload the model."""
    try:
        generator = get_generator()
        success = generator.load_model()
        return jsonify({
            'success': success,
            'device': generator.device,
            'mode': 'template' if generator.use_template_mode else 'model'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/ai-generate-questions', methods=['POST'])
def api_ai_generate_questions():
    """Generate fresh practice questions using the fine-tuned AI model."""
    import random
    from question_paper_generator import NCERT_KNOWLEDGE
    
    data = request.get_json()
    chapter = data.get('chapter', '')
    question_type = data.get('type', 'mcq')  # mcq, short_answer, concept
    count = min(int(data.get('count', 3)), 5)  # Max 5 at a time
    lang = data.get('lang', 'en')
    
    if not chapter:
        return jsonify({'success': False, 'error': 'Chapter name is required'}), 400
    
    results = []
    
    # Try AI model first
    model, tokenizer = get_ai_model()
    if model is not None:
        for i in range(count):
            if question_type == 'mcq':
                prompt = f"### Instruction:\nGenerate an MCQ question about '{chapter}'.\n\n### Input:\n\n### Response:"
            elif question_type == 'short_answer':
                prompt = f"### Instruction:\nGenerate a short answer question about '{chapter}'.\n\n### Input:\n\n### Response:"
            else:  # concept
                prompt = f"### Instruction:\nExplain a key concept from the chapter '{chapter}'.\n\n### Input:\n\n### Response:"
            
            response = generate_ai_text(prompt, max_new_tokens=200)
            if response:
                if question_type == 'mcq' and 'Question:' in response:
                    # Parse MCQ format
                    try:
                        lines = response.strip().split('\n')
                        q_text = ''
                        options = []
                        answer = ''
                        explanation = ''
                        for line in lines:
                            line = line.strip()
                            if line.startswith('Question:'):
                                q_text = line.replace('Question:', '').strip()
                            elif line.startswith(('A)', 'B)', 'C)', 'D)')):
                                options.append(line[2:].strip())
                            elif line.startswith('Answer:'):
                                answer = line.replace('Answer:', '').strip()
                            elif line.startswith('Explanation:'):
                                explanation = line.replace('Explanation:', '').strip()
                        
                        if q_text and len(options) >= 4 and answer:
                            results.append({
                                'type': 'mcq',
                                'question': q_text,
                                'options': options[:4],
                                'answer': answer,
                                'explanation': explanation,
                                'source': 'ai_generated'
                            })
                    except:
                        pass
                elif question_type == 'short_answer' and 'Question:' in response:
                    try:
                        parts = response.split('Answer:')
                        q_text = parts[0].replace('Question:', '').strip()
                        ans_text = parts[1].strip() if len(parts) > 1 else ''
                        if q_text and ans_text:
                            results.append({
                                'type': 'short_answer',
                                'question': q_text,
                                'answer': ans_text,
                                'source': 'ai_generated'
                            })
                    except:
                        pass
                else:
                    results.append({
                        'type': 'concept',
                        'question': f'AI explanation about {chapter}',
                        'answer': response,
                        'source': 'ai_generated'
                    })
    
    # If AI didn't produce enough, supplement from knowledge bank
    knowledge = NCERT_KNOWLEDGE.get(chapter, {})
    if len(results) < count and knowledge:
        if question_type == 'mcq':
            pool = list(knowledge.get('mcq_pool', []))
            random.shuffle(pool)
            for q, opts, ans, exp in pool[:count - len(results)]:
                results.append({
                    'type': 'mcq',
                    'question': q,
                    'options': opts,
                    'answer': ans,
                    'explanation': exp,
                    'source': 'knowledge_bank'
                })
        elif question_type == 'short_answer':
            pool = list(knowledge.get('short_answers', []))
            random.shuffle(pool)
            for q, ans in pool[:count - len(results)]:
                results.append({
                    'type': 'short_answer',
                    'question': q,
                    'answer': ans,
                    'source': 'knowledge_bank'
                })
        else:
            pool = list(knowledge.get('key_concepts', []))
            random.shuffle(pool)
            for term, defn in pool[:count - len(results)]:
                results.append({
                    'type': 'concept',
                    'question': f'Explain: {term}',
                    'answer': defn,
                    'source': 'knowledge_bank'
                })
    
    # Translate if needed
    translate_questions_bulk(results, lang)

    return jsonify({
        'success': True,
        'questions': results,
        'count': len(results),
        'ai_used': model is not None
    })


@app.route('/api/ai-model-status', methods=['GET'])
def api_ai_model_status():
    """Check if the fine-tuned AI model is available."""
    use_hf_api = os.environ.get('USE_HF_API', 'False').lower() == 'true'
    model_exists = use_hf_api or (FINETUNED_MODEL_PATH.exists() and (FINETUNED_MODEL_PATH / "adapter_config.json").exists())
    return jsonify({
        'success': True,
        'model_available': model_exists,
        'model_loaded': _ai_model is not None,
        'model_path': str(FINETUNED_MODEL_PATH)
    })


@app.route('/api/health', methods=['GET'])
def api_health():
    """Health check endpoint."""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'topics_available': len(get_available_topics())
    })


@app.route('/api/export-docx', methods=['POST'])
def api_export_docx():
    """Export question paper as Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        
        data = request.get_json()
        if not data or 'paper' not in data:
            return jsonify({'success': False, 'error': 'No paper data provided'}), 400
        
        paper = data['paper']
        
        # Create Word document
        doc = Document()
        
        # Set up styles
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        # Add Page Borders
        sec = doc.sections[0]
        sectPr = sec._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # border thickness
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), 'auto')
            pgBorders.append(border)
        sectPr.append(pgBorders)
        
        # Add header table for logo, "The Joy of being", and JGS text
        header_table = doc.add_table(rows=1, cols=3)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Col 1: Logo and AP018
        cell_logo = header_table.cell(0, 0)
        cell_logo.width = Inches(2.0)
        logo_path = os.path.join(app.root_path, 'static', 'jgs_logo.png')
        if os.path.exists(logo_path):
            para = cell_logo.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run()
            run.add_picture(logo_path, width=Inches(1.5))
            para2 = cell_logo.add_paragraph()
            para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = para2.add_run("AP018")
            run2.bold = True
            run2.font.size = Pt(10)
            
        # Col 2: The Joy of being
        cell_center = header_table.cell(0, 1)
        cell_center.width = Inches(2.0)
        para_center = cell_center.paragraphs[0]
        para_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_center = para_center.add_run("\n\nThe Joy of being")
        run_center.bold = True
        run_center.font.size = Pt(10)

        # Col 3: Exam Code
        cell_right = header_table.cell(0, 2)
        cell_right.width = Inches(2.0)
        para_right = cell_right.paragraphs[0]
        para_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_right = para_right.add_run("JGS/EXAM/IF-02/R01")
        run_right.bold = True
        run_right.font.size = Pt(10)
        
        # Add exam title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(12)
        title.paragraph_format.space_after = Pt(12)
        run = title.add_run("ANNUAL EXAMINATION [2025-2026]")
        run.bold = True
        run.font.size = Pt(14)
        
        # Metadata table
        meta_table = doc.add_table(rows=3, cols=6)
        meta_table.style = 'Table Grid'
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Get topic safely
        topics_meta = paper.get('topicMetadata', [])
        topic_name = topics_meta[0].get('chapter', 'Various') if topics_meta else 'Various Topics'
        subject_name = topics_meta[0].get('subject', 'Science') if topics_meta else 'Science'
        class_name = topics_meta[0].get('class', 'X') if topics_meta else 'X'
        
        # Row 1: Merged
        cell_name = meta_table.cell(0, 0)
        cell_name.merge(meta_table.cell(0, 5))
        para_name = cell_name.paragraphs[0]
        run_name = para_name.add_run("NAME OF THE STUDENT:")
        run_name.bold = True
        run_name.italic = True
        
        # Row 2: Headers
        headers = ["CL / SEC", "SUBJECT", "DATE", "TIME", "MARKS", "PAGE"]
        for i, header_text in enumerate(headers):
            cell = meta_table.cell(1, i)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(header_text)
            run.bold = True
            run.italic = True
            
        # Row 3: Values
        values = [class_name, subject_name.upper(), "25.02.2026", paper.get('duration', '2 HRS.'), str(paper.get('totalMarks', 80)), "1 of 4"]
        for i, val in enumerate(values):
            cell = meta_table.cell(2, i)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(val)
            run.bold = True

        # Instructions
        inst_para = doc.add_paragraph()
        inst_para.paragraph_format.space_before = Pt(12)
        inst_para.paragraph_format.space_after = Pt(12)
        inst_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_inst = inst_para.add_run("Instructions: Answers to this paper must be written on the paper provided separately. You will not be allowed to write during the first 15 minutes. This time is to be spent in reading the question paper. The time given at the head of the paper is the time allowed for writing the answers. Marks will be deducted if questions or bits of questions are numbered incorrectly. The intended marks for questions or parts of questions are given in brackets ( ).")
        run_inst.bold = True
        run_inst.italic = True
        

        
        # Add sections
        for section in paper.get('sections', []):
            
            # --- Format New Section Header ---
            sec_header = doc.add_paragraph()
            sec_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sec_header.paragraph_format.space_before = Pt(18)
            run = sec_header.add_run(f"SECTION - {section.get('name', 'A')}")
            run.bold = True
            run.underline = True
            
            # --- Format Section Instructions ---
            is_compulsory = section.get('isCompulsory', True)
            attempt_cnt = section.get('attemptCount', len(section.get('questions', [])))
            q_cnt = len(section.get('questions', []))
            
            attempt_text = "(Attempt all questions from this section)" if is_compulsory or attempt_cnt >= q_cnt else f"(Attempt any {attempt_cnt} questions from this section)"
            marks_cnt_line = f"({q_cnt if is_compulsory else attempt_cnt}x{section.get('marksPerQuestion', 1)}={section.get('totalMarks', 0)}M)"
            
            sec_inst = doc.add_paragraph()
            sec_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sec_inst.paragraph_format.space_after = Pt(12)
            run_inst2 = sec_inst.add_run(attempt_text)
            run_inst2.bold = True
            run_inst2.italic = True
            
            # --- Format Question block headers ---
            q_type = section.get('questionType', 'mcq')
            q_instruction = "Choose the correct answers to the questions from the given options."
            if q_type == 'fill_blank': q_instruction = "Fill in the blanks with appropriate words."
            elif q_type in ('very_short', 'short'): q_instruction = "Answer the following questions briefly."
            elif q_type == 'long': q_instruction = "Answer the following questions in detail."
            
            # Question 1 Label
            q_label = doc.add_paragraph()
            run_label = q_label.add_run("Question 1")
            run_label.bold = True
            
            # Instruction and marks line (align left, right)
            inst_marks_table = doc.add_table(rows=1, cols=2)
            inst_marks_table.autofit = False
            cell_left = inst_marks_table.cell(0, 0)
            cell_right = inst_marks_table.cell(0, 1)
            cell_left.width = Inches(5.0)
            cell_right.width = Inches(1.5)
            
            p_left = cell_left.paragraphs[0]
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r_left = p_left.add_run(q_instruction)
            r_left.bold = True
            
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_right = p_right.add_run(marks_cnt_line)
            r_right.bold = True
            
            doc.add_paragraph() # Spacing
            
            # Questions in section
            for i, q in enumerate(section.get('questions', [])):
                q_para = doc.add_paragraph()
                
                # We use the current index instead of the raw number so it resets per section if needed.
                q_text = q.get('question', q.get('text', ''))
                
                # Check for image URL
                image_url = q.get('imageUrl')
                if image_url:
                    try:
                        import requests
                        from io import BytesIO
                        # Download image
                        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
                        response = requests.get(image_url, headers=headers)
                        if response.status_code == 200:
                            image_stream = BytesIO(response.content)
                            # Add paragraph for image, centered
                            img_para = doc.add_paragraph()
                            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = img_para.add_run()
                            # Resize to fit neatly (e.g., 2.5 inches max width) to save space
                            run.add_picture(image_stream, width=Inches(2.5))
                    except Exception as e:
                        print(f"Error adding image to docx: {e}")
                
                q_para.add_run(f"{i+1}. ").bold = True
                q_para.add_run(f"{q_text}")
                
                # Add options for MCQ if it exists
                if 'options' in q and isinstance(q['options'], list):
                    import re
                    for idx, opt in enumerate(q['options']):
                        letter = chr(97 + idx) # a, b, c, d
                        # Strip existing 'A) ', 'b. ', '(C) ', '1. ', etc.
                        clean_opt = re.sub(r'^[\(\[]?[a-dA-D1-4][\)\].]\s*', '', str(opt).strip())
                        doc.add_paragraph(f"    ({letter}) {clean_opt}")
                        
            doc.add_paragraph()  # Spacing between sections
            
        # Save to BytesIO
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        filename = f"question_paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except ImportError:
        return jsonify({'success': False, 'error': 'python-docx not installed. Run: pip install python-docx'}), 500
    except Exception as e:
        print(f"Error exporting to Word: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/export-lesson-plan-docx', methods=['POST'])
def api_export_lesson_plan_docx():
    """Export lesson plan as Word document in school format."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        data = request.get_json()
        if not data or 'lessonPlan' not in data:
            return jsonify({'success': False, 'error': 'No lesson plan data provided'}), 400
        
        plan = data['lessonPlan']
        lang = data.get('lang', 'en')
        
        # Define language labels
        LP_LABELS = {
            'en': {
                'prereq': 'PRE-REQUISITE KNOWLEDGE:',
                'details': 'LESSON PLAN DETAILS:',
                'duration': 'Duration', 'outcome': 'Learning Outcome', 'methodology': 'Methodology/Strategy',
                'aids': 'Teaching Aids', 'teacher_act': 'Teacher Activities', 'learner_act': 'Learner Activities',
                'assessment': 'Assessment', 'homework': 'Home Assignment',
                'values': 'VALUES DEALT IN THE LESSON (VD): ',
                'reallife': 'REAL LIFE APPLICATION (RL): ',
                'crosscurr': 'CROSS CURRICULAR CONNECTION (CC): ', 'extended': 'EXTENDED TASK: ',
                'sig_teacher': 'Signature of Subject Teacher', 'sig_hod': 'Signature of HOD', 'sig_principal': 'Signature of Principal',
                'lbl_teacher': 'Teacher:', 'lbl_grade': 'Grade:', 'lbl_subject': 'Subject:',
                'lbl_periods': 'Periods:', 'lbl_lesson': 'Lesson:', 'lbl_plan': 'Plan No:', 'lbl_topic': 'Topic:',
                'lbl_class': 'Class', 'lbl_lesson_plan': 'LESSON PLAN'
            },
            'te': {
                'prereq': 'అవసరమైన అర్హత జ్ఞానం:',
                'details': 'పాఠ్య ప్రణాళిక వివరాలు:',
                'duration': 'వ్యవధి', 'outcome': 'అభ్యాస ఫలితం', 'methodology': 'బోధన పద్ధతి/వ్యూహం',
                'aids': 'బోధన సహాయాలు', 'teacher_act': 'ఉపాధ్యాయుని కార్యకలాపాలు', 'learner_act': 'విద్యార్థి కార్యకలాపాలు',
                'assessment': 'మూల్యాంకనం', 'homework': 'గృహపాఠం',
                'values': 'పాఠంలో చర్చించిన విలువలు: ', 'reallife': 'నిజ జీవిత అనువర్తనం: ',
                'crosscurr': 'అంతర్-విషయ సంబంధం: ', 'extended': 'విస్తరించిన పని: ',
                'sig_teacher': 'సబ్జెక్ట్ ఉపాధ్యాయుని సంతకం', 'sig_hod': 'HOD సంతకం', 'sig_principal': 'ప్రిన్సిపల్ సంతకం',
                'lbl_teacher': 'ఉపాధ్యాయుడు:', 'lbl_grade': 'తరగతి:', 'lbl_subject': 'సబ్జెక్ట్:',
                'lbl_periods': 'పిరియడ్లు:', 'lbl_lesson': 'పాఠం:', 'lbl_plan': 'ప్రణాళిక నం:', 'lbl_topic': 'అంశం:',
                'lbl_class': 'తరగతి', 'lbl_lesson_plan': 'పాఠ్య ప్రణాళిక'
            },
            'hi': {
                'prereq': 'पूर्व-आवश्यक ज्ञान:',
                'details': 'पाठ योजना विवरण:',
                'duration': 'अवधि', 'outcome': 'सीखने का परिणाम', 'methodology': 'शिक्षण पद्धति/रणनीति',
                'aids': 'शिक्षण सामग्री', 'teacher_act': 'शिक्षक गतिविधियाँ', 'learner_act': 'छात्र गतिविधियाँ',
                'assessment': 'मूल्यांकन', 'homework': 'गृहकार्य',
                'values': 'पाठ में संबोधित मूल्य: ', 'reallife': 'वास्तविक जीवन अनुप्रयोग: ',
                'crosscurr': 'अंतर-विषय संबंध: ', 'extended': 'विस्तारित कार्य: ',
                'sig_teacher': 'विषय शिक्षक के हस्ताक्षर', 'sig_hod': 'HOD के हस्ताक्षर', 'sig_principal': 'प्रधानाचार्य के हस्ताक्षर',
                'lbl_teacher': 'शिक्षक:', 'lbl_grade': 'कक्षा:', 'lbl_subject': 'विषय:',
                'lbl_periods': 'कालांश:', 'lbl_lesson': 'पाठ:', 'lbl_plan': 'योजना क्र.:', 'lbl_topic': 'विषय:',
                'lbl_class': 'कक्षा', 'lbl_lesson_plan': 'पाठ योजना'
            },
            'mr': {
                'prereq': 'पूर्व-ज्ञान:',
                'details': 'पाठ योजना तपशील:',
                'duration': 'कालावधी', 'outcome': 'शिक्षण परिणाम', 'methodology': 'शिक्षण पद्धती/रणनीती',
                'aids': 'शिक्षण साधने', 'teacher_act': 'शिक्षक उपक्रम', 'learner_act': 'विद्यार्थी उपक्रम',
                'assessment': 'मूल्यमापन', 'homework': 'गृहपाठ',
                'values': 'पाठातील मूल्ये: ', 'reallife': 'वास्तव जीवनातील उपयोजन: ',
                'crosscurr': 'आंतर-विषय संबंध: ', 'extended': 'विस्तारित कार्य: ',
                'sig_teacher': 'विषय शिक्षकाची सही', 'sig_hod': 'HOD ची सही', 'sig_principal': 'मुख्याध्यापकाची सही',
                'lbl_teacher': 'शिक्षक:', 'lbl_grade': 'इयत्ता:', 'lbl_subject': 'विषय:',
                'lbl_periods': 'तास:', 'lbl_lesson': 'पाठ:', 'lbl_plan': 'योजना क्र.:', 'lbl_topic': 'विषय:',
                'lbl_class': 'इयत्ता', 'lbl_lesson_plan': 'पाठ योजना'
            },
            'ta': {
                'prereq': 'முன் தேவையான அறிவு:',
                'details': 'பாட திட்ட விவரங்கள்:',
                'duration': 'கால அளவு', 'outcome': 'கற்றல் விளைவு', 'methodology': 'கற்பித்தல் முறை/உத்தி',
                'aids': 'கற்பித்தல் உதவிகள்', 'teacher_act': 'ஆசிரியர் செயல்பாடுகள்', 'learner_act': 'மாணவர் செயல்பாடுகள்',
                'assessment': 'மதிப்பீடு', 'homework': 'வீட்டுப்பாடம்',
                'values': 'பாடத்தில் கூறப்பட்ட மதிப்புகள்: ', 'reallife': 'நிஜ வாழ்க்கை பயன்பாடு: ',
                'crosscurr': 'பாட இணைப்புகள்: ', 'extended': 'நீட்டிக்கப்பட்ட பணி: ',
                'sig_teacher': 'பாட ஆசிரியர் கையொப்பம்', 'sig_hod': 'HOD கையொப்பம்', 'sig_principal': 'அதிபர் கையொப்பம்',
                'lbl_teacher': 'ஆசிரியர்:', 'lbl_grade': 'வகுப்பு:', 'lbl_subject': 'பாடம்:',
                'lbl_periods': 'பாட நேரம்:', 'lbl_lesson': 'பாடம்:', 'lbl_plan': 'திட்ட எண்:', 'lbl_topic': 'தலைப்பு:',
                'lbl_class': 'வகுப்பு', 'lbl_lesson_plan': 'பாட திட்டம்'
            },
            'kn': {
                'prereq': 'ಪೂರ್ವ-ಅಗತ್ಯ ಜ್ಞಾನ:',
                'details': 'ಪಾಠ ಯೋಜನೆ ವಿವರಗಳು:',
                'duration': 'ಅವಧಿ', 'outcome': 'ಕಲಿಕಾ ಫಲಿತಾಂಶ', 'methodology': 'ಬೋಧನಾ ವಿಧಾನ/ತಂತ್ರ',
                'aids': 'ಬೋಧನಾ ಸಾಧನಗಳು', 'teacher_act': 'ಶಿಕ್ಷಕ ಚಟುವಟಿಕೆಗಳು', 'learner_act': 'ವಿದ್ಯಾರ್ಥಿ ಚಟುವಟಿಕೆಗಳು',
                'assessment': 'ಮೌಲ್ಯಮಾಪನ', 'homework': 'ಗೃಹಕಾರ್ಯ',
                'values': 'ಪಾಠದಲ್ಲಿ ಸಂಬೋಧಿಸಿದ ಮೌಲ್ಯಗಳು: ', 'reallife': 'ನಿಜ ಜೀವನ ಅನ್ವಯ: ',
                'crosscurr': 'ಅಂತರ ಪಠ್ಯ ಸಂಪರ್ಕ: ', 'extended': 'ವಿಸ್ತರಿತ ಕಾರ್ಯ: ',
                'sig_teacher': 'ವಿಷಯ ಶಿಕ್ಷಕರ ಸಹಿ', 'sig_hod': 'HOD ಸಹಿ', 'sig_principal': 'ಪ್ರಾಂಶುಪಾಲರ ಸಹಿ',
                'lbl_teacher': 'ಶಿಕ್ಷಕ:', 'lbl_grade': 'ತರಗತಿ:', 'lbl_subject': 'ವಿಷಯ:',
                'lbl_periods': 'ಅವಧಿಗಳು:', 'lbl_lesson': 'ಪಾಠ:', 'lbl_plan': 'ಯೋಜನೆ ಸಂ.:', 'lbl_topic': 'ವಿಷಯ:',
                'lbl_class': 'ತರಗತಿ', 'lbl_lesson_plan': 'ಪಾಠ ಯೋಜನೆ'
            }
        }
        
        lbl = LP_LABELS.get(lang, LP_LABELS['en'])
        
        # Create Word document
        doc = Document()
        
        # Set up styles
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        
        # ===== HEADER =====
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(plan.get('schoolName', 'School Name'))
        run.bold = True
        run.font.size = Pt(16)
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(lbl['lbl_lesson_plan'])
        run.bold = True
        run.font.size = Pt(14)
        
        # ===== INFO TABLE =====
        info_table = doc.add_table(rows=4, cols=4)
        info_table.style = 'Table Grid'
        
        # Row 1
        info_table.cell(0, 0).text = lbl['lbl_teacher']
        info_table.cell(0, 1).text = plan.get('teacherName', '')
        info_table.cell(0, 2).text = 'Date:' # not translatable directly unless we added it, but let's keep it Date for now, wait we can add to dict or just let it translate in deep_translator if needed, actually it wasn't requested
        info_table.cell(0, 3).text = plan.get('date', '')
        
        # Row 2
        info_table.cell(1, 0).text = lbl['lbl_subject']
        info_table.cell(1, 1).text = plan.get('subject', '')
        info_table.cell(1, 2).text = lbl['lbl_grade']
        info_table.cell(1, 3).text = f"{lbl['lbl_class']} {plan.get('grade', '')}"
        
        # Row 3
        info_table.cell(2, 0).text = lbl['lbl_lesson']
        info_table.cell(2, 1).text = plan.get('chapterName', '')
        info_table.cell(2, 2).text = lbl['lbl_periods']
        info_table.cell(2, 3).text = f"{plan.get('periodsCount', 1)} ({plan.get('periodDuration', 40)} min each)"
        
        # Row 4
        info_table.cell(3, 0).text = lbl['lbl_topic']
        info_table.cell(3, 1).merge(info_table.cell(3, 3))
        info_table.cell(3, 1).text = plan.get('topic', '')
        
        # Bold the labels
        for row in info_table.rows:
            for i in [0, 2]:
                if i < len(row.cells):
                    for para in row.cells[i].paragraphs:
                        for run in para.runs:
                            run.bold = True
        
        doc.add_paragraph()  # Spacing
        
        # ===== PRE-REQUISITE KNOWLEDGE =====
        prereq_header = doc.add_paragraph()
        run = prereq_header.add_run(lbl['prereq'])
        run.bold = True
        run.font.size = Pt(11)
        
        for item in plan.get('prerequisiteKnowledge', []):
            doc.add_paragraph(f"• {item}", style='List Bullet')
        
        doc.add_paragraph()  # Spacing
        
        # ===== MAIN LESSON PLAN TABLE =====
        main_header = doc.add_paragraph()
        run = main_header.add_run(lbl['details'])
        run.bold = True
        run.font.size = Pt(11)
        
        # Create table with 8 columns
        phases = plan.get('phases', [])
        table = doc.add_table(rows=len(phases) + 1, cols=8)
        table.style = 'Table Grid'
        
        # Header row
        headers = [lbl['duration'], lbl['outcome'], lbl['methodology'], lbl['aids'], 
                   lbl['teacher_act'], lbl['learner_act'], lbl['assessment'], lbl['homework']]
        for i, header_text in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header_text
            # Make header bold
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        
        # Data rows
        for row_idx, phase in enumerate(phases, 1):
            table.cell(row_idx, 0).text = f"{phase.get('name', '')}\n({phase.get('duration', '')})"
            table.cell(row_idx, 1).text = phase.get('learningOutcome', '')
            table.cell(row_idx, 2).text = ', '.join(phase.get('methodology', []))
            table.cell(row_idx, 3).text = ', '.join(phase.get('teachingAids', []))
            table.cell(row_idx, 4).text = '\n'.join([f"• {a}" for a in phase.get('teacherActivities', [])])
            table.cell(row_idx, 5).text = '\n'.join([f"• {a}" for a in phase.get('learnerActivities', [])])
            table.cell(row_idx, 6).text = phase.get('assessment', '')
            table.cell(row_idx, 7).text = phase.get('homeAssignment', '-')
        
        # Set font size for all cells
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.space_after = Pt(2)
                    for run in para.runs:
                        run.font.size = Pt(8)
        
        doc.add_paragraph()  # Spacing
        
        # ===== FOOTER SECTIONS =====
        # Values
        values_para = doc.add_paragraph()
        run = values_para.add_run(lbl['values'])
        run.bold = True
        values_para.add_run(plan.get('values', ''))
        
        # Real Life Application
        rl_para = doc.add_paragraph()
        run = rl_para.add_run(lbl['reallife'])
        run.bold = True
        rl_para.add_run(plan.get('realLifeApplication', ''))
        
        # Cross Curricular
        cc_para = doc.add_paragraph()
        run = cc_para.add_run(lbl['crosscurr'])
        run.bold = True
        cc_para.add_run(plan.get('crossCurricular', ''))
        
        # Extended Task
        if plan.get('extendedTask'):
            ext_para = doc.add_paragraph()
            run = ext_para.add_run(lbl['extended'])
            run.bold = True
            ext_para.add_run(plan.get('extendedTask', ''))
        
        doc.add_paragraph()  # Spacing
        doc.add_paragraph()  # More spacing
        
        # ===== SIGNATURES =====
        sig_table = doc.add_table(rows=2, cols=3)
        sig_table.cell(0, 0).text = ''
        sig_table.cell(0, 1).text = ''
        sig_table.cell(0, 2).text = ''
        
        sig_table.cell(1, 0).text = f"________________________\n{lbl['sig_teacher']}"
        sig_table.cell(1, 1).text = f"________________________\n{lbl['sig_hod']}"
        sig_table.cell(1, 2).text = f"________________________\n{lbl['sig_principal']}"
        
        for cell in sig_table.rows[1].cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to BytesIO
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Generate filename
        chapter_safe = ''.join(c if c.isalnum() else '_' for c in plan.get('chapterName', 'lesson'))
        filename = f"lesson_plan_{chapter_safe}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except ImportError:
        return jsonify({'success': False, 'error': 'python-docx not installed. Run: pip install python-docx'}), 500
    except Exception as e:
        print(f"Error exporting lesson plan to Word: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/start-finetuning', methods=['POST'])
def api_start_finetuning():
    """Trigger the fine-tuning process in the background."""
    global _training_status
    
    if _training_status["is_training"]:
        return jsonify({"success": False, "error": "Training is already in progress."})
        
    def run_training():
        global _training_status, _ai_model_loaded
        _training_status.update({
            "is_training": True,
            "current_step": "generating_data",
            "progress": 10,
            "last_error": None,
            "start_time": datetime.now().isoformat(),
            "logs": ["Starting data generation..."]
        })
        
        try:
            # 1. Generate new training data (includes dynamic uploads)
            from train_pipeline import generate_training_data
            generate_training_data()
            _training_status["logs"].append("Data generation complete.")
            _training_status["current_step"] = "training_model"
            _training_status["progress"] = 30
            
            # 2. Run training via subprocess to avoid blocking/memory issues
            import os
            venv_python = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), '.venv', 'Scripts', 'python.exe')
            if not os.path.exists(venv_python):
                venv_python = sys.executable
                
            process = subprocess.Popen(
                [venv_python, "train_pipeline.py", "train"],
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                encoding='utf-8',  # Explicitly use utf-8 for Windows compatibility
                errors='replace',  # Replace undecodable characters instead of crashing
                bufsize=1,
                creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == 'win32' else 0
            )
            
            _training_status["logs"].append("Model training started...")
            
            # Use safe iterator for stdout to catch potential pipe errors
            try:
                for line in iter(process.stdout.readline, ""):
                    line = line.strip()
                    if line:
                        _training_status["logs"].append(line)
                        
                        # Special parsing for our custom ProgressPrinterCallback
                        if "PROGRESS_UPDATE" in line:
                            # Format: PROGRESS_UPDATE | Epoch: 1.23 | Step: 100 | Loss: 0.85
                            try:
                                # Echo to terminal as per user request
                                print(f"[FINE-TUNE] {line}")
                                
                                # Estimate progress (approx 2000 steps for 5 epochs)
                                parts = line.split('|')
                                step_str = [p for p in parts if "Step:" in p][0]
                                current_step = int(step_str.split(':')[1].strip())
                                # Total steps estimation: 5 epochs * (1600 samples / 4 grad_acc) = 2000
                                _training_status["progress"] = min(99, (current_step / 2000) * 100)
                            except:
                                pass
                        else:
                            # Standard output echo
                            print(f"[TRAINING] {line}")
                            
                            # Fallback progress estimation
                            if "%" in line:
                                _training_status["progress"] = min(95, _training_status["progress"] + 0.1)
            except Exception as pipe_err:
                _training_status["logs"].append(f"Pipe log error (non-fatal): {str(pipe_err)}")
                
            process.wait()
            
            if process.returncode == 0:
                _training_status["logs"].append("Training successfully finished!")
                _training_status["progress"] = 100
                _training_status["current_step"] = "completed"
                # Flag to reload model on next inference
                _ai_model_loaded = False
            else:
                _training_status["logs"].append(f"Training failed with exit code {process.returncode}")
                _training_status["current_step"] = "failed"
                _training_status["last_error"] = f"Subprocess exit code: {process.returncode}"
                
        except Exception as e:
            _training_status["last_error"] = str(e)
            _training_status["current_step"] = "failed"
            _training_status["logs"].append(f"CRITICAL ERROR: {str(e)}")
        finally:
            _training_status["is_training"] = False

    thread = threading.Thread(target=run_training)
    thread.daemon = True
    thread.start()
    
    return jsonify({"success": True, "message": "Training started in background."})

@app.route('/api/training-status')
def api_training_status():
    """Get the current training status and latest logs."""
    status_copy = _training_status.copy()
    # Only send latest logs to keep response small
    status_copy["logs"] = _training_status["logs"][-20:]
    return jsonify(status_copy)


if __name__ == '__main__':
    print("=" * 60)
    print("VedLinks AI Educational Content Generator")
    print("=" * 60)
    print(f"\n[Directory] Topics directory: {DATA_TOPICS_PATH.absolute()}")
    
    topics = get_available_topics()
    print(f"[Info] Available topics: {len(topics)}")
    
    for topic in topics[:5]:  # Show first 5
        print(f"   - Class {topic['class']} {topic['subject']}: {topic['chapter']}")
    
    if len(topics) > 5:
        print(f"   ... and {len(topics) - 5} more")
    
    import torch
    cuda_status = "[OK] CUDA Available" if torch.cuda.is_available() else "[WARN] CUDA Not Found (Using CPU)"
    gpu_name = torch.cuda.get_device_name(0) if torch.cuda.is_available() else "N/A"
    
    print(f"\n[System] Hardware Status: {cuda_status}")
    if torch.cuda.is_available():
        print(f"[System] GPU: {gpu_name}")
        
    print("\n[Server] Starting server at http://127.0.0.1:5000")
    print("\n[API] API Endpoints:")
    print("   GET  /api/topics          - List all topics")
    print("   GET  /api/topic/<id>      - Get topic details")
    print("   POST /api/generate-paper  - Generate question paper")
    print("   POST /api/generate-lesson-plan - Generate lesson plan")
    print("=" * 60)
    
    app.run(debug=True, use_reloader=False, host='127.0.0.1', port=5000)
