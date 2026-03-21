import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_detailed_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('VedLinks: Comprehensive Project Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Overview
    doc.add_heading('1. Executive Overview', level=1)
    doc.add_paragraph(
        "VedLinks is an advanced AI-powered educational platform designed to transform traditional "
        "textbooks into structured, interactive, and adaptive learning experiences. Designed tightly "
        "around the NCERT curriculum, it serves two main audiences:\n"
        "1. Educators: Offering high-quality, reproducible assessment and planning tools.\n"
        "2. Students: Providing an adaptive, personalized AI tutor with contextual verification."
    )
    doc.add_paragraph(
        "The platform is driven by a custom-trained Small Language Model (SLM), based on TinyLlama-1.1B-Chat, "
        "which is fine-tuned strictly on factual NCERT textbook embeddings."
    )

    # Architecture
    doc.add_heading('2. Project Architecture & Stack', level=1)
    architecture_text = (
        "• Backend Framework: Python with Flask (app.py), serving robust JSON APIs.\n"
        "• AI Processing: Hugging Face transformers ecosystem (LoRA/QLoRA) for inference and local fine-tuning.\n"
        "• Data Persistence: JSON-based local document stores (data/topics/, data/topic_registry.json) instead of complex databases, "
        "ensuring high portability and speed.\n"
        "• Frontend: Vanilla HTML/CSS/JS (templates and static directories) rendering interactive widgets and dynamic loaders.\n"
        "• Model Output: Saved in output/qlora_tuned_model/."
    )
    doc.add_paragraph(architecture_text)

    # Core Features
    doc.add_heading('3. Core Features & Capabilities', level=1)

    doc.add_heading('Feature 1: Smart Question Paper Generator', level=2)
    doc.add_paragraph(
        "This is the most heavy-weight curriculum tool available to educators, managed by question_paper_generator.py.\n"
        "• Multi-Format Generation: Mathematically balances and generates four formats of questions based on a selected template: "
        "MCQs, Fill-in-the-Blanks, Short Answer, and Long Answer.\n"
        "• Difficulty Controlling: Internally maps difficulties (Easy, Medium, Hard) guaranteeing exams are balanced.\n"
        "• Professional Exporting: Generates an automated, fully styled, ready-to-print Microsoft Word (.docx) file "
        "containing the full paper and a separate Answer Key/Marking Scheme at the end.\n"
        "• Current Coverage: Includes structured textbook data explicitly seeded for Class 6 and Class 8 Science chapters."
    )

    doc.add_heading('Feature 2: Educator Lesson Planner', level=2)
    doc.add_paragraph(
        "Managed by src/lesson_planner.py and accessed via /api/generate-lesson-plan.\n"
        "• Automated Structuring: Dynamically compiles a structured curriculum guide per chapter suitable for school administration records.\n"
        "• Exporting: Capable of saving the generated lesson plan into a localized Word Document using python-docx."
    )

    doc.add_heading('Feature 3: Textbook Onboarding & Processing', level=2)
    doc.add_paragraph(
        "The pipeline that converts a static PDF into an interactive AI dataset.\n"
        "• Smart Metadata Labelling: Users upload raw PDF files which are immediately classified by Class, Subject, and Chapter via /api/upload-textbook.\n"
        "• Automatic Indexing: Strips text out of the PDF, cleans it, segments the topics, and writes it directly into the data/topics/ directory "
        "as structured JSON metadata, simultaneously updating the master data/topic_registry.json. Managed by src/pdf_processor.py."
    )

    doc.add_heading('Feature 4: Adaptive AI Practice & Doubts (The Student End)', level=2)
    doc.add_paragraph(
        "Found at the /practice frontend route and backed by /api/practice-questions.\n"
        "• Topic Contextualization: Retrieves real knowledge base concepts from JSON via /api/concepts so practice stays factual.\n"
        "• Hybrid Serving: Combines strictly verified 'NCERT Knowledge Bank' questions (badged 📚) with AI-generated questions (badged 🤖).\n"
        "• On-Demand AI Generation: If static questions are exhausted, students can click 'Generate More'. The Flask backend uses the TinyLlama-1.1B LoRA model "
        "to synthesize completely fresh, contextual MCQs on the fly via /api/ai-generate-questions.\n"
        "• Instant Feedback: UI evaluates selection immediately, tracking scores and providing explanations."
    )

    # AI Pipeline
    doc.add_heading('4. The AI Fine-Tuning Pipeline', level=1)
    doc.add_paragraph(
        "A uniquely powerful aspect of VedLinks is its local fine-tuning orchestrator, designed to prevent the baseline AI from hallucinating un-related facts.\n"
        "• Target Model: TinyLlama-1.1B-Chat-v1.0.\n"
        "• Dataset Pipeline: Generates a JSONL repository (data/finetune_dataset.jsonl with 1,172+ curated instructional pairs) formatting the model to act as an expert NCERT tutor. Managed by src/dataset_generator.py.\n"
        "• Background Orchestrator: Web admins trigger a QLoRA fine-tuning session (/api/start-finetuning). This spawns a background thread that executes the deep-learning loop (src/train_lora.py).\n"
        "• Live Monitoring: Fetches log streams and step percentages from the background thread to show real-time progress on the UI via /api/training-status.\n"
        "• GPU Optimizations: Specific configuration scripts (check_gpu.py, enable_gpu.py) aimed at Windows CUDA capabilities to democratize training on desktop GPUs."
    )

    # API Suite
    doc.add_heading('5. Elite API Suite Mapping', level=1)
    doc.add_paragraph(
        "The complete API ecosystem built on Flask:\n"
        "1. /api/upload-textbook (POST): Extracts textual data from user uploads, builds JSON registry.\n"
        "2. /api/get-topics (GET): Returns a flat list of all parsed and available topics.\n"
        "3. /api/get-topics-grouped (GET): Returns a hierarchical tree (Class -> Subject -> Chapter).\n"
        "4. /api/concepts (POST): Retrieves textbook factual summaries for the 'Doubts' feature.\n"
        "5. /api/practice-questions (POST): Dispatches a batch of mixed Bank + AI questions for grading.\n"
        "6. /api/generate-paper (POST): Invokes the structured test formulation logic.\n"
        "7. /api/export-docx (POST): Converts the last generated exam paper into a binary download stream.\n"
        "8. /api/export-lesson-plan-docx (POST): Exports lesson plans as .docx files.\n"
        "9. /api/ai-generate-questions (POST): Directly inferences the LoRA weights to create spontaneous MCQs.\n"
        "10. /api/start-finetuning (POST): Initiates the training loop background worker.\n"
        "11. /api/training-status (GET): Polling endpoint for the active training loop.\n"
        "12. /api/model-status (GET) & /api/ai-model-status (GET): Reports model availability."
    )

    # Codebase Structure
    doc.add_heading('6. Codebase File Structure', level=1)
    doc.add_paragraph(
        "App Logic:\n"
        "• app.py: Core Application Server and routing.\n"
        "• question_paper_generator.py: Heavy lifting for test paper creation.\n"
        "• train_pipeline.py & train_simple.py: Orchestration for dataset creation and logic tying.\n"
        "\nSource Code (src/):\n"
        "• dataset_generator.py, lesson_planner.py, pdf_processor.py, train_lora.py.\n"
        "\nData (data/):\n"
        "• Stores unstructured (raw PDFs) and structured (topics/ JSONs). Finetuning dataset generated here.\n"
        "\nWeb Interface:\n"
        "• templates/: HTML for index.html, practice.html, lesson_planner.html, upload.html.\n"
        "• static/: JS/CSS handling interactions (e.g. lesson_planner.js).\n"
        "\nDeployment:\n"
        "• vercel.json, Dockerfile, hf_space_deploy/ for versatile hosting choices."
    )

    # Output to File
    output_path = "VedLinks_Detailed_Project_Report.docx"
    doc.save(output_path)
    print(f"Report successfully saved to {output_path}")

if __name__ == '__main__':
    create_detailed_report()
